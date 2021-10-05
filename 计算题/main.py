import random, os
import PySimpleGUI as sg
from docx import Document
from docx.shared import RGBColor, Pt, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

'''
该程序产生口算题doc文件：

'''
document = Document()

'''
生成随机数，默认产生2位数的随机数
'''


def randomtoNUM(bit=2):
    if bit == 2:
        return random.randint(10, 99)
    elif bit == 3:
        return random.randint(100, 999)
    elif bit == 1:
        return random.randint(1, 9)


'''
2位数减法(含退位)
'''


def chutijian():
    jia1_10 = random.randint(1, 9)
    jia1_1 = random.randint(0, 8)
    jia1 = int(str(jia1_10) + str(jia1_1))
    # print(8-int(str(jia1)[0]))
    jia2_max10 = int(str(jia1)[0]) - 1
    if jia2_max10 < 1:
        jia2_10 = jia2_max10
    else:
        jia2_10 = random.randint(0, jia2_max10)

    if int(str(jia1)[1]) + 1 >= 9:
        jia2_1 = 9
    else:
        jia2_1 = random.randint(int(str(jia1)[1]) + 1, 9)
    jia2 = int(str(jia2_10) + str(jia2_1))
    # print(eval("{} + {}".format(jia1,jia2)))

    return "{:<2d} － {:>2d}".format(jia1, jia2)


'''
2位数加法(含进位)
'''


def chutiAdd():
    jia1_10 = random.randint(1, 8)
    jia1_1 = randomtoNUM(1)
    jia1 = int(str(jia1_10) + str(jia1_1))
    # print(8-int(str(jia1)[0]))
    jia2_max10 = 8 - int(str(jia1)[0])
    if jia2_max10 <= 1:
        jia2_10 = abs(jia2_max10)
    else:
        jia2_10 = random.randint(1, jia2_max10)

    if 10 - int(str(jia1)[1]) >= 9:
        jia2_1 = 9
    else:
        jia2_1 = random.randint(10 - int(str(jia1)[1]), 9)
    jia2 = int(str(jia2_10) + str(jia2_1))
    # print(eval("{} + {}".format(jia1,jia2)))
    if random.randint(0, 1):
        return "{:<2d} ＋ {:>2d}".format(jia1, jia2)
    else:
        return "{:<2d} ＋ {:>2d}".format(jia2, jia1)


'''
1位数加法(含进位)
'''


def chuti1BitAddH():
    jia1 = randomtoNUM(1)
    jia2 = 10 - jia1  # 第二个数最小值

    if jia2 == 1:
        jia2 = 9
    else:
        jia2 = random.randint(10 - jia1, 9)

    if random.randint(0, 1):
        return "{:<2d} ＋ {:>2d}".format(jia1, jia2)
    else:
        return "{:<2d} ＋ {:>2d}".format(jia2, jia1)


'''
1位数加法(不含进位)
'''


def chuti1BitAdd():
    jia1 = random.randint(1, 8)
    jia2 = 9 - jia1  # 第二个数最大值

    if jia2 <= 1:
        jia2 = 1
    else:
        jia2 = random.randint(1, jia2)

    if random.randint(0, 1):
        return "{:<2d} ＋ {:>2d}".format(jia1, jia2)
    else:
        return "{:<2d} ＋ {:>2d}".format(jia2, jia1)


'''
1位数减法
'''


def chuti1Bitsub():
    jian1 = random.randint(2, 9)
    jian2 = jian1 - 1

    if jian2 <= 1:
        jian2 = 1
    else:
        jian2 = random.randint(1, jian2)

    return "{:<2d} － {:>2d}".format(jian1, jian2)

'''
产生2个相减需退位的数，即第1个数小于第二个数
'''
def gtTenSub():
    jia1 = random.randint(0, 8)
    jia2 = jia1+1

    if jia2 >= 9:
        jia2 = 9
    else:
        jia2 = random.randint(jia2, 9)

    return jia1, jia2

'''
产生2个相减不退位的数，第1个数大于第二个数
'''


def ltTenSub():
    jia1 = random.randint(2, 9)
    jia2 = jia1-1  # 第二个数最大值

    if jia2 <= 1:
        jia2 = 1
    else:
        jia2 = random.randint(1, jia2)
    return jia1, jia2



'''
产生2个相加进位的数
'''


def gtTen():
    jia1 = random.randint(1, 9)
    jia2 = 9 - jia1

    if jia2 >= 9:
        jia2 = 9
    else:
        jia2 = random.randint(jia2 + 1, 9)

    return jia1, jia2


'''
产生2个相加不进位的数
'''


def ltTen():
    jia1 = random.randint(1, 8)
    jia2 = 9 - jia1  # 第二个数最大值

    if jia2 <= 1:
        jia2 = 1
    else:
        jia2 = random.randint(1, jia2)
    return jia1, jia2


'''
3位数相加

'''


def chuti3BitAdd():
    listone = []
    listtwo = []
    gttenNum = random.randint(1, 3)  # 选出需要进位的个数
    for i in range(gttenNum):
        one, two = gtTen()
        listone.append(str(one))
        listtwo.append(str(two))

    for j in range(3 - gttenNum):
        one, two = ltTen()
        listone.append(str(one))
        listtwo.append(str(two))
    #print('完整：', ''.join(listone), ''.join(listtwo))
    if random.randint(0, 1):
        return "{:>8s} \n＋ {:>5s}\n￣￣￣￣￣\n\n".format(''.join(listone), ''.join(listtwo))
    else:
        return "{:>8s} \n＋ {:>5s}\n￣￣￣￣￣\n\n".format(''.join(listtwo), ''.join(listone))



'''
3位数相减

'''


def chuti3BitSub():
    listone = []
    listtwo = []
    gttenNum = random.randint(1, 2)  # 选出需要退位的个数
    for i in range(gttenNum):
        one, two = gtTenSub()
        listone.append(str(one))
        listtwo.append(str(two))

    for j in range(3 - gttenNum):
        one, two = ltTenSub()
        listone.append(str(one))
        listtwo.append(str(two))
    #print('完整：', ''.join(listone), ''.join(listtwo))
    listone.reverse()
    listtwo.reverse()
    return "{:>8s} \n－ {:>5d}\n￣￣￣￣￣\n\n".format(''.join(listone), int(''.join(listtwo)))


'''

整十或整百加减
'''


def allplugsub():
    localhandleFunctionlist = ['chuti1BitAdd', 'chuti1BitAddH', 'chuti1Bitsub', 'chutiAdd', 'chutijian']
    resulttemp = eval(random.choice(localhandleFunctionlist) + "()")
    strarr = resulttemp.split('－')
    isFlagplugsub = len(strarr)  # 是1为加法，2为减法
    beishu = random.choice(['0', '00'])
    if isFlagplugsub == 2:  # 减法
        jian1 = strarr[0].strip() + beishu
        jian2 = strarr[1].strip() + beishu
        return "{:<2d} － {:>2d}".format(int(jian1), int(jian2))
    else:  # 加法
        strarr = resulttemp.split('＋')
        jia1 = strarr[0].strip() + beishu
        jia2 = strarr[1].strip() + beishu
        if random.randint(0, 1):
            return "{:<2d} ＋ {:>2d}".format(int(jia1), int(jia2))
        else:
            return "{:<2d} ＋ {:>2d}".format(int(jia2), int(jia1))


'''
产生doc文档
'''


def makedoc(handlelist,pagenum):
    '''
    p = document.add_paragraph()

    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(u'小学生计算题练习 ')
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.font.size = Pt(28)
    p.space_before = Pt(40)

    p.add_run().add_break(break_type=6)
    '''

    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(16)

    ti = document.add_paragraph()
    ti.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for nump in range(pagenum):
        tishu = 45
        if len(handlelist) ==1 and ('chuti3BitAdd' in handlelist or 'chuti3BitSub' in handlelist):
            tishu = 21
        if len(handlelist) ==2 and ('chuti3BitAdd' in handlelist and 'chuti3BitSub' in handlelist):
            tishu = 21

        for i in range(tishu):  # 产生45题口算题

            # 根据函数名称随机选择题目类型，handlelist list的值，可用的函数'chutiAdd', 'chutijian','chuti1BitAddH','chuti1BitAdd'
            if tishu == 45:
                content = u"{:<7}=".format(eval(random.choice(handlelist) + "()"))
                run1 = ti.add_run("{:<40}".format(content))
            elif tishu == 21:
                content = u"{}".format(eval(random.choice(handlelist) + "()"))
                if i==20 and nump==pagenum-1:
                    content=content.strip('\n\n')

                run1 = ti.add_run("{}".format(content))
            run1.font.color.rgb = RGBColor(0, 0, 0)
            run1.font.size = Pt(16)
            if tishu == 45:
                ti.paragraph_format.line_spacing = Pt(45)

    #print(len(document.sections))
    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = '                     小学生计算题练习                  '
        paragraph.style = document.styles["Header"]

        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

        # Columns = 1
        sectPr = section._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'), '3')

    document.save('test.docx')  # 可以设置其他路径

    os.startfile("test.docx")


# ff     GreenTan
class win():
    sg.ChangeLookAndFeel('LightGreen')
    ok_btn = sg.SimpleButton('生成习题', size=(10, 2), font=("微软雅黑", 12), button_color=('white', 'firebrick3'))
    cancel_btn = sg.Button('关闭程序', size=(10, 2), font=("微软雅黑", 12))
    layout = [
        [sg.Checkbox('1位数加，不进位', default=False, size=(35, 2), font=("微软雅黑", 12))],
        [sg.Checkbox('1位数加，进位', size=(35, 2), font=("微软雅黑", 12))],
        [sg.Checkbox('1位数减', size=(35, 2), font=("微软雅黑", 12))],
        [sg.Checkbox('2位数加，进位', size=(35, 2), font=("微软雅黑", 12))],
        [sg.Checkbox('2位数减，退位', size=(35, 2), font=("微软雅黑", 12))],
        [sg.Checkbox('整十或整百加减', size=(35, 2), font=("微软雅黑", 12))],
        [sg.Checkbox('3位数相加,进位,只能与竖式同时选择', size=(35, 2),text_color=('blue'), font=("微软雅黑", 12))],
        [sg.Checkbox('3位数减法,退位,只能与竖式同时选择', size=(35, 2),text_color=('blue'), font=("微软雅黑", 12))],
    # [sg.Slider(range=(1,5), orientation='h',size=(35, 10),  font=("微软雅黑", 12))],
        #[sg.Text('This is some text', font='Courier 12', text_color='blue', background_color='green')],
        #[sg.Spin([1,2,3,4,5], size=(35, 10),  font=("微软雅黑", 12))],
    [sg.Text('要生成几页：',auto_size_text=True,size=(15, 1), font=("微软雅黑", 12))],
    [sg.Slider(range=(1,6), orientation='h',size=(35, 15),  font=("微软雅黑", 12))],
        [ok_btn, cancel_btn],[sg.StatusBar('好好学习，天天向上\n谢修一，加油',size=(400,10), font=("微软雅黑", 12))]

    ]

    window = sg.Window('谢修一数字学习系统', default_element_size=(40, 2), size=(400, 485)).Layout(layout)

    handleFunctionlist = ['chuti1BitAdd', 'chuti1BitAddH', 'chuti1Bitsub', 'chutiAdd', 'chutijian', 'allplugsub',
                        'chuti3BitAdd','chuti3BitSub']
    while True:
        event, values = window.read()
        if event in (None, '关闭程序'):
            # User closed the Window or hit the Cancel button
            break
        elif event in (None, '生成习题'):
            # sg.Popup(event, values)
            handlelist = []
            #print(values)

            for i in values:

                if type(values[i])==bool and values[i]:
                    handlelist.append(handleFunctionlist[i])
                #print(handlelist)
            pagenum=int(values[8])
            #print(pagenum)
            makedoc(handlelist,pagenum)

            break

if __name__ == '__main__':
    win()
    